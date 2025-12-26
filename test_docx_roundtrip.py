#!/usr/bin/env python3
"""
Test 1: DOCX Round-Trip Comment Insertion

This script takes a .docx file exported from Google Docs, adds a comment
at a specific text location, and saves a new file for re-upload.

Usage:
    pip install python-docx
    python test_docx_roundtrip.py test_input.docx test_output.docx

The output file can then be uploaded to Google Drive (with conversion to Google Docs)
to test if the comment position is preserved.
"""

import sys
from docx import Document
from docx.shared import Inches


def find_and_comment_text(doc: Document, target_text: str, comment_text: str, author: str = "Test Author") -> bool:
    """
    Find target_text in the document and add a comment to it.

    Returns True if the text was found and commented, False otherwise.
    """
    for paragraph in doc.paragraphs:
        if target_text in paragraph.text:
            # We need to find the run(s) containing the target text
            # This is tricky because text might span multiple runs

            # First, let's try a simple approach: if the paragraph has runs,
            # find which run contains our target text
            for run in paragraph.runs:
                if target_text in run.text:
                    # Found it in a single run - this is the easy case
                    comment = doc.add_comment(
                        runs=[run],
                        text=comment_text,
                        author=author,
                        initials=author[:2].upper()
                    )
                    print(f"✓ Added comment to text: '{target_text}'")
                    print(f"  Comment: '{comment_text}'")
                    return True

            # If we get here, the target text spans multiple runs
            # For this test, we'll comment the whole paragraph instead
            if paragraph.runs:
                comment = doc.add_comment(
                    runs=paragraph.runs,
                    text=comment_text,
                    author=author,
                    initials=author[:2].upper()
                )
                print(f"✓ Added comment to paragraph containing: '{target_text}'")
                print(f"  (Target text spanned multiple runs, commented whole paragraph)")
                print(f"  Comment: '{comment_text}'")
                return True

    print(f"✗ Could not find text: '{target_text}'")
    return False


def main():
    if len(sys.argv) < 3:
        print("Usage: python test_docx_roundtrip.py <input.docx> <output.docx>")
        print("\nExample:")
        print("  python test_docx_roundtrip.py test_input.docx test_output.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    print(f"\n=== DOCX Round-Trip Comment Test ===\n")
    print(f"Input:  {input_file}")
    print(f"Output: {output_file}")
    print()

    # Load the document
    try:
        doc = Document(input_file)
    except Exception as e:
        print(f"Error loading document: {e}")
        sys.exit(1)

    # Print document structure for debugging
    print("Document structure:")
    for i, para in enumerate(doc.paragraphs):
        preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
        print(f"  Paragraph {i}: '{preview}' ({len(para.runs)} runs)")
    print()

    # Add a comment to our target phrase
    target = "This is a target phrase for our comment test"
    comment = "TEST COMMENT: This comment was added programmatically via python-docx"

    success = find_and_comment_text(doc, target, comment)

    if not success:
        print("\nFallback: Adding comment to first non-empty paragraph...")
        for para in doc.paragraphs:
            if para.text.strip() and para.runs:
                doc.add_comment(
                    runs=para.runs,
                    text=comment,
                    author="Test Author",
                    initials="TA"
                )
                print(f"✓ Added comment to: '{para.text[:50]}...'")
                break

    # Save the modified document
    doc.save(output_file)
    print(f"\n✓ Saved: {output_file}")

    print("\n=== Next Steps ===")
    print("1. Upload test_output.docx to Google Drive")
    print("2. Right-click > 'Open with' > 'Google Docs'")
    print("   (This converts it to Google Docs format)")
    print("3. Check if the comment appears anchored to the correct text")
    print("4. Report back: Is the comment attached to 'This is a target phrase...'?")


if __name__ == "__main__":
    main()
