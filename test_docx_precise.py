#!/usr/bin/env python3
"""
Test 1b: DOCX Round-Trip with Precise Comment Targeting

This improved script splits runs to isolate the target text,
allowing comments to be anchored to specific phrases rather than whole paragraphs.

Usage:
    pip install python-docx
    python test_docx_precise.py test_input.docx test_output_precise.docx

"""

import sys
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def split_run_at_text(paragraph, target_text: str):
    """
    Find target_text in the paragraph and split runs so it's isolated.

    Returns a list of runs containing exactly the target text, or None if not found.
    """
    # First, build a map of character positions to runs
    full_text = ""
    run_boundaries = []  # [(start_pos, end_pos, run_index), ...]

    for i, run in enumerate(paragraph.runs):
        start = len(full_text)
        full_text += run.text
        end = len(full_text)
        run_boundaries.append((start, end, i))

    # Find target in full text
    target_start = full_text.find(target_text)
    if target_start == -1:
        return None

    target_end = target_start + len(target_text)

    print(f"  Found '{target_text}' at positions {target_start}-{target_end}")
    print(f"  Full paragraph text: '{full_text}'")
    print(f"  Run boundaries: {run_boundaries}")

    # Find which runs contain our target
    runs_to_split = []
    for start, end, run_idx in run_boundaries:
        # Check if this run overlaps with our target
        if start < target_end and end > target_start:
            runs_to_split.append({
                'run_idx': run_idx,
                'run_start': start,
                'run_end': end,
                'target_start_in_run': max(0, target_start - start),
                'target_end_in_run': min(end - start, target_end - start)
            })

    print(f"  Runs to modify: {runs_to_split}")

    # Now we need to split the runs
    # This is complex because we need to manipulate the underlying XML
    target_runs = []

    for split_info in runs_to_split:
        run_idx = split_info['run_idx']
        run = paragraph.runs[run_idx]
        run_text = run.text

        t_start = split_info['target_start_in_run']
        t_end = split_info['target_end_in_run']

        before_text = run_text[:t_start]
        target_part = run_text[t_start:t_end]
        after_text = run_text[t_end:]

        print(f"  Run {run_idx}: before='{before_text}', target='{target_part}', after='{after_text}'")

        if before_text == "" and after_text == "":
            # The whole run is part of our target
            target_runs.append(run)
        else:
            # We need to split this run
            # For now, let's try a simpler approach: just modify the existing run
            # and note what we need

            # Store original formatting
            run_element = run._element

            if before_text:
                # Create a new run for the "before" text
                new_run_before = copy.deepcopy(run_element)
                for t in new_run_before.findall(qn('w:t')):
                    t.text = before_text
                run_element.addprevious(new_run_before)

            if after_text:
                # Create a new run for the "after" text
                new_run_after = copy.deepcopy(run_element)
                for t in new_run_after.findall(qn('w:t')):
                    t.text = after_text
                run_element.addnext(new_run_after)

            # Modify the current run to contain only the target part
            for t in run_element.findall(qn('w:t')):
                t.text = target_part

            target_runs.append(run)

    return target_runs


def add_comment_to_text(doc: Document, target_text: str, comment_text: str, author: str = "Test Author") -> bool:
    """
    Find target_text and add a comment to exactly that text.
    """
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if target_text in paragraph.text:
            print(f"\nProcessing paragraph {para_idx}...")

            target_runs = split_run_at_text(paragraph, target_text)

            if target_runs:
                comment = doc.add_comment(
                    runs=target_runs,
                    text=comment_text,
                    author=author,
                    initials=author[:2].upper()
                )
                print(f"\n✓ Added comment to exactly: '{target_text}'")
                return True
            else:
                print(f"✗ Could not isolate target text in runs")
                return False

    print(f"✗ Target text not found in any paragraph")
    return False


def main():
    if len(sys.argv) < 3:
        print("Usage: python test_docx_precise.py <input.docx> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    print(f"\n=== DOCX Precise Comment Test ===\n")
    print(f"Input:  {input_file}")
    print(f"Output: {output_file}")

    doc = Document(input_file)

    # Show document structure
    print("\nDocument structure:")
    for i, para in enumerate(doc.paragraphs):
        print(f"  Para {i}: '{para.text[:60]}...' ({len(para.runs)} runs)")
        for j, run in enumerate(para.runs):
            print(f"    Run {j}: '{run.text}'")

    # Target the specific phrase
    target = "This is a target phrase for our comment test"
    comment = "PRECISE TEST: This should highlight ONLY the target phrase"

    success = add_comment_to_text(doc, target, comment)

    if success:
        doc.save(output_file)
        print(f"\n✓ Saved: {output_file}")
        print("\n=== Next Steps ===")
        print("1. Upload to Google Drive")
        print("2. Open with Google Docs")
        print("3. Check if comment highlights ONLY the target phrase")
    else:
        print("\n✗ Failed to add precise comment")
        sys.exit(1)


if __name__ == "__main__":
    main()
