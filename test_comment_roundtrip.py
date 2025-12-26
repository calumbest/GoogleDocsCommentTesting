#!/usr/bin/env python3
"""
Test 2: Comment Preservation Round-Trip

Tests whether existing comments survive when we add new comments via python-docx.

Usage:
    python test_comment_roundtrip.py test_with_comment.docx test_roundtrip_output.docx

Steps:
1. Create a Google Doc with manual comment(s)
2. Export as .docx (e.g., test_with_comment.docx)
3. Run this script
4. Upload output to Google Drive
5. Verify both old and new comments appear
"""

import sys
import copy
from docx import Document
from docx.oxml.ns import qn


def list_existing_comments(doc: Document):
    """List all existing comments in the document."""
    print("\n=== Existing Comments ===")

    # Access comments through the document's part
    try:
        comments_part = doc.part.comments_part
        if comments_part is None:
            print("No comments found in document.")
            return []

        comments_element = comments_part.element
        comments = comments_element.findall(qn('w:comment'))

        if not comments:
            print("No comments found in document.")
            return []

        print(f"Found {len(comments)} comment(s):\n")

        comment_list = []
        for comment in comments:
            comment_id = comment.get(qn('w:id'))
            author = comment.get(qn('w:author'))
            date = comment.get(qn('w:date'))

            # Get comment text
            text_elements = comment.findall('.//' + qn('w:t'))
            text = ''.join(t.text for t in text_elements if t.text)

            print(f"  Comment ID: {comment_id}")
            print(f"  Author: {author}")
            print(f"  Date: {date}")
            print(f"  Text: '{text}'")
            print()

            comment_list.append({
                'id': comment_id,
                'author': author,
                'date': date,
                'text': text
            })

        return comment_list

    except Exception as e:
        print(f"Error reading comments: {e}")
        return []


def find_comment_anchors(doc: Document):
    """Find where comments are anchored in the document."""
    print("\n=== Comment Anchors in Document ===")

    for para_idx, para in enumerate(doc.paragraphs):
        para_xml = para._element

        # Look for commentRangeStart elements
        range_starts = para_xml.findall('.//' + qn('w:commentRangeStart'))
        range_ends = para_xml.findall('.//' + qn('w:commentRangeEnd'))

        if range_starts or range_ends:
            print(f"\nParagraph {para_idx}: '{para.text[:50]}...'")
            for rs in range_starts:
                print(f"  commentRangeStart id={rs.get(qn('w:id'))}")
            for re in range_ends:
                print(f"  commentRangeEnd id={re.get(qn('w:id'))}")


def split_run_at_text(paragraph, target_text: str):
    """Find target_text and split runs to isolate it."""
    full_text = ""
    run_boundaries = []

    for i, run in enumerate(paragraph.runs):
        start = len(full_text)
        full_text += run.text
        end = len(full_text)
        run_boundaries.append((start, end, i))

    target_start = full_text.find(target_text)
    if target_start == -1:
        return None

    target_end = target_start + len(target_text)

    runs_to_split = []
    for start, end, run_idx in run_boundaries:
        if start < target_end and end > target_start:
            runs_to_split.append({
                'run_idx': run_idx,
                'run_start': start,
                'run_end': end,
                'target_start_in_run': max(0, target_start - start),
                'target_end_in_run': min(end - start, target_end - start)
            })

    target_runs = []

    for split_info in runs_to_split:
        run_idx = split_info['run_idx']
        run = paragraph.runs[run_idx]
        run_text = run.text

        t_start = split_info['target_start_in_run']
        t_end = split_info['target_end_in_run']

        before_text = run_text[:t_start]
        target_part = run_text[t_start:t_end]
        after_text = run_text[t_end:]

        if before_text == "" and after_text == "":
            target_runs.append(run)
        else:
            run_element = run._element

            if before_text:
                new_run_before = copy.deepcopy(run_element)
                for t in new_run_before.findall(qn('w:t')):
                    t.text = before_text
                run_element.addprevious(new_run_before)

            if after_text:
                new_run_after = copy.deepcopy(run_element)
                for t in new_run_after.findall(qn('w:t')):
                    t.text = after_text
                run_element.addnext(new_run_after)

            for t in run_element.findall(qn('w:t')):
                t.text = target_part

            target_runs.append(run)

    return target_runs


def add_new_comment(doc: Document, target_text: str, comment_text: str, author: str = "Python Script"):
    """Add a new comment to target text."""
    for paragraph in doc.paragraphs:
        if target_text in paragraph.text:
            target_runs = split_run_at_text(paragraph, target_text)
            if target_runs:
                doc.add_comment(
                    runs=target_runs,
                    text=comment_text,
                    author=author,
                    initials=author[:2].upper()
                )
                return True
    return False


def main():
    if len(sys.argv) < 3:
        print("Usage: python test_comment_roundtrip.py <input.docx> <output.docx>")
        print("\nThis script tests whether existing comments survive when adding new ones.")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    print(f"\n{'='*50}")
    print("Test 2: Comment Preservation Round-Trip")
    print(f"{'='*50}")
    print(f"\nInput:  {input_file}")
    print(f"Output: {output_file}")

    # Load document
    doc = Document(input_file)

    # Show document structure
    print("\n=== Document Structure ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            preview = para.text[:60] + "..." if len(para.text) > 60 else para.text
            print(f"  Para {i}: '{preview}'")

    # List existing comments
    existing_comments = list_existing_comments(doc)

    # Show where comments are anchored
    find_comment_anchors(doc)

    # Add a new comment
    print("\n=== Adding New Comment ===")
    target = "quick brown fox"
    new_comment = "NEW COMMENT: Added programmatically - testing preservation"

    if add_new_comment(doc, target, new_comment):
        print(f"✓ Added new comment to: '{target}'")
    else:
        # Try alternate target
        target = "Lorem ipsum"
        if add_new_comment(doc, target, new_comment):
            print(f"✓ Added new comment to: '{target}'")
        else:
            print("✗ Could not find target text for new comment")
            print("  Will save anyway to test existing comment preservation")

    # Save
    doc.save(output_file)
    print(f"\n✓ Saved: {output_file}")

    # Verify output
    print("\n=== Verifying Output File ===")
    doc_verify = Document(output_file)
    final_comments = list_existing_comments(doc_verify)

    print(f"\n{'='*50}")
    print("SUMMARY")
    print(f"{'='*50}")
    print(f"Existing comments in input:  {len(existing_comments)}")
    print(f"Total comments in output:    {len(final_comments)}")
    print(f"New comments added:          {len(final_comments) - len(existing_comments)}")

    print("\n=== Next Steps ===")
    print("1. Upload output file to Google Drive")
    print("2. Open with Google Docs")
    print("3. Verify:")
    print("   - Original comment(s) still present and anchored?")
    print("   - New comment present and anchored to 'quick brown fox'?")


if __name__ == "__main__":
    main()
